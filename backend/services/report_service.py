"""
Report Service

Handles generation of Word documents from material analysis results.
"""

from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from typing import Dict, Any
import io
import logging

from backend.utils.logger import get_logger

logger = get_logger(__name__)


class ReportService:
    """Service for generating Word reports from analysis results."""

    def __init__(self):
        self.logger = logger

    def generate_docx_report(self, analysis_result: Dict[str, Any]) -> bytes:
        """
        Generate Word document from analysis results.

        Args:
            analysis_result: Analysis result from data_analyzer_service

        Returns:
            bytes: Word document as bytes
        """
        try:
            self.logger.info("Generating Word report")

            # Create new document
            doc = Document()

            # Add title
            title = doc.add_heading('Material Analysis Report', 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Add analysis summary
            self._add_summary_section(doc, analysis_result)

            # Add technique-specific sections
            features = analysis_result.get('features', {})
            for technique, data in features.items():
                if 'error' not in data:
                    self._add_technique_section(doc, technique, data)

            # Add LLM analysis section
            summary = analysis_result.get('summary', '')
            if summary:
                self._add_llm_analysis_section(doc, summary)

            # Add recommendations section
            self._add_recommendations_section(doc, analysis_result)

            # Save to bytes buffer
            buffer = io.BytesIO()
            doc.save(buffer)
            buffer.seek(0)

            self.logger.info("Word report generated successfully")
            return buffer.getvalue()

        except Exception as e:
            self.logger.error(f"Error generating Word report: {str(e)}")
            raise ValueError(f"Failed to generate report: {str(e)}")

    def _add_summary_section(self, doc: Document, analysis_result: Dict[str, Any]):
        """Add summary section to document."""
        doc.add_heading('Analysis Summary', level=1)

        features = analysis_result.get('features', {})
        techniques_analyzed = [tech for tech, data in features.items() if 'error' not in data]

        if techniques_analyzed:
            p = doc.add_paragraph()
            p.add_run(f"Analysis completed for the following techniques: {', '.join(techniques_analyzed)}")
        else:
            p = doc.add_paragraph()
            p.add_run("No successful analyses completed.")

    def _add_technique_section(self, doc: Document, technique: str, data: Dict[str, Any]):
        """Add technique-specific section to document."""
        doc.add_heading(f'{technique.upper()} Analysis', level=1)

        # Add basic information
        if 'source_file' in data:
            p = doc.add_paragraph()
            p.add_run(f"Source file: {data['source_file']}")

        # Add technique-specific data
        if technique == 'xrd':
            self._add_xrd_data(doc, data)
        elif technique == 'ir':
            self._add_ir_data(doc, data)
        elif technique == 'tga':
            self._add_tga_data(doc, data)
        elif technique == 'bet':
            self._add_bet_data(doc, data)

    def _add_xrd_data(self, doc: Document, data: Dict[str, Any]):
        """Add XRD-specific data to document."""
        doc.add_heading('XRD Results', level=2)

        # Create table for key metrics
        table = doc.add_table(rows=1, cols=2)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Parameter'
        hdr_cells[1].text = 'Value'

        # Add data rows
        metrics = [
            ('Peak Count', data.get('peak_count', 'N/A')),
            ('Significant Peaks', data.get('significant_peaks', 'N/A')),
            ('Max Intensity', data.get('max_intensity', 'N/A')),
            ('Max Intensity 2θ', data.get('max_intensity_theta', 'N/A')),
            ('Data Points', data.get('data_points', 'N/A'))
        ]

        for param, value in metrics:
            row_cells = table.add_row().cells
            row_cells[0].text = param
            row_cells[1].text = str(value)

        # Add peaks information
        peaks = data.get('peaks', [])
        if peaks:
            doc.add_heading('Key Peaks', level=3)
            p = doc.add_paragraph()
            p.add_run("Top peaks identified:")
            for i, (theta, intensity) in enumerate(peaks[:5]):  # Top 5 peaks
                doc.add_paragraph(f"{i+1}. 2θ = {theta:.2f}°, Intensity = {intensity:.2f}")

    def _add_ir_data(self, doc: Document, data: Dict[str, Any]):
        """Add IR-specific data to document."""
        doc.add_heading('IR Results', level=2)

        # Create table for key metrics
        table = doc.add_table(rows=1, cols=2)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Parameter'
        hdr_cells[1].text = 'Value'

        # Add data rows
        metrics = [
            ('Peak Count', data.get('peak_count', 'N/A')),
            ('Data Points', data.get('data_points', 'N/A')),
            ('Max Intensity', data.get('max_intensity', 'N/A')),
            ('Min Intensity', data.get('min_intensity', 'N/A'))
        ]

        for param, value in metrics:
            row_cells = table.add_row().cells
            row_cells[0].text = param
            row_cells[1].text = str(value)

        # Add functional groups
        functional_groups = data.get('functional_groups', {})
        if functional_groups:
            doc.add_heading('Functional Groups Detected', level=3)
            for group, info in functional_groups.items():
                doc.add_paragraph(f"{group}: Peak at {info.get('peak_wavenumber', 'N/A')} cm⁻¹")

    def _add_tga_data(self, doc: Document, data: Dict[str, Any]):
        """Add TGA-specific data to document."""
        doc.add_heading('TGA Results', level=2)

        # Create table for key metrics
        table = doc.add_table(rows=1, cols=2)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Parameter'
        hdr_cells[1].text = 'Value'

        # Add data rows
        metrics = [
            ('Sample Name', data.get('sample_name', 'Unknown')),
            ('Unit Adsorption', data.get('unit_adsorption', 'N/A')),
            ('Desorption Energy', data.get('desorption_energy', 'N/A')),
            ('Raw Data Available', 'Yes' if data.get('raw_data_available') else 'No')
        ]

        for param, value in metrics:
            row_cells = table.add_row().cells
            row_cells[0].text = param
            row_cells[1].text = str(value)

    def _add_bet_data(self, doc: Document, data: Dict[str, Any]):
        """Add BET-specific data to document."""
        doc.add_heading('BET Results', level=2)

        # Create table for key metrics
        table = doc.add_table(rows=1, cols=2)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Parameter'
        hdr_cells[1].text = 'Value'

        # Add data rows
        metrics = [
            ('Surface Area', f"{data.get('surface_area', 'N/A')} m²/g"),
            ('Pore Size', f"{data.get('pore_size', 'N/A')} nm"),
            ('Pore Volume', f"{data.get('pore_volume', 'N/A')} cm³/g"),
            ('Source File', data.get('source_file', 'Unknown'))
        ]

        for param, value in metrics:
            row_cells = table.add_row().cells
            row_cells[0].text = param
            row_cells[1].text = str(value)

    def _add_llm_analysis_section(self, doc: Document, summary: str):
        """Add LLM analysis section to document."""
        doc.add_heading('AI Analysis', level=1)

        # Add summary text
        p = doc.add_paragraph()
        p.add_run(summary)

    def _add_recommendations_section(self, doc: Document, analysis_result: Dict[str, Any]):
        """Add recommendations section to document."""
        doc.add_heading('Recommendations', level=1)

        # Extract recommendations from LLM analysis if available
        summary = analysis_result.get('summary', '')
        if 'recommendations' in summary.lower():
            # Try to extract recommendations from summary
            lines = summary.split('\n')
            for line in lines:
                if line.strip().startswith('-') or line.strip().startswith('•'):
                    doc.add_paragraph(line.strip())
        else:
            doc.add_paragraph("Based on the analysis results, consider the following:")
            doc.add_paragraph("• Review the material characterization data for quality assurance")
            doc.add_paragraph("• Compare results with literature values for validation")
            doc.add_paragraph("• Consider additional characterization techniques if needed")


def create_report_service() -> ReportService:
    """Factory function to create report service instance."""
    return ReportService()
